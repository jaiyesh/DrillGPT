import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from drilling.utils import print_app_banner, logger
from drilling.data import generate_synthetic_drilling_data
from drilling.models import DrillingModels
from drilling.optimizer import ParameterOptimizer
from drilling.llm import LLMService
from drilling.assistant import DrillingAssistant
from colorama import Fore, Style
import time
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    hr.append(bottom)
    p._p.get_or_add_pPr().append(hr)

def add_markdown_runs(para, text):
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    tokens = re.findall(pattern, text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = para.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = para.add_run(token[1:-1])
            run.italic = True
        else:
            para.add_run(token)

def format_markdown_to_word(doc, text):
    lines = text.split('\n')
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('### '):
            para = doc.add_paragraph(stripped[4:], style='Heading 3')
        elif stripped.startswith('## '):
            para = doc.add_paragraph(stripped[3:], style='Heading 2')
        elif stripped.startswith('# '):
            para = doc.add_paragraph(stripped[2:], style='Heading 1')
        elif stripped.startswith('---') or stripped.startswith('==='):
            add_horizontal_line(doc)
        elif re.match(r'^\d+\.', stripped):
            para = doc.add_paragraph(style='List Number')
            add_markdown_runs(para, stripped)
        elif re.match(r'^[-*•]\s+', stripped):
            para = doc.add_paragraph(style='List Bullet')
            content = re.sub(r'^[-*•]\s+', '', stripped)
            add_markdown_runs(para, content)
        elif '**' in stripped or '*' in stripped:
            para = doc.add_paragraph()
            add_markdown_runs(para, stripped)
        else:
            doc.add_paragraph(line)

def clean_text(text):
    ansi_escape = re.compile(r'\x1B[@-_][0-?]*[ -/]*[@-~]')
    text = ansi_escape.sub('', text)
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    return text

def run_drilling_optimization_demo():
    print_app_banner()
    logger.info(Fore.GREEN + "1. Generating synthetic drilling data..." + Style.RESET_ALL)
    print("\n1. Generating synthetic drilling data...")
    drilling_data = generate_synthetic_drilling_data(n_samples=500)
    logger.info(Fore.CYAN + f"   Generated {len(drilling_data)} data points" + Style.RESET_ALL)
    print(f"   Generated {len(drilling_data)} data points")

    logger.info(Fore.GREEN + "2. Training ML models..." + Style.RESET_ALL)
    print("\n2. Training ML models...")
    models = DrillingModels()
    models.train_rop_model(drilling_data)
    models.train_bit_failure_model(drilling_data)

    logger.info(Fore.GREEN + "3. Setting up parameter optimizer and drilling assistant..." + Style.RESET_ALL)
    print("\n3. Setting up parameter optimizer and drilling assistant...")
    optimizer = ParameterOptimizer(models)
    llm_service = LLMService()
    assistant = DrillingAssistant(llm_service, models)

    logger.info(Fore.GREEN + "4. Simulating real-time drilling optimization..." + Style.RESET_ALL)
    print("\n4. Simulating real-time drilling optimization...")
    sim_indices = [100, 200, 300]

    for i, idx in enumerate(sim_indices):
        logger.info(Fore.MAGENTA + f"\n--- Step {i+1}/{len(sim_indices)} ---" + Style.RESET_ALL)
        print(f"\n--- Step {i+1}/{len(sim_indices)} ---")
        current_state = drilling_data.iloc[idx].to_dict()
        assistant.current_state = current_state
        logger.info(Fore.YELLOW + f"Depth: {current_state['depth']:.1f} ft, Formation: {current_state['formation']}" + Style.RESET_ALL)
        print(f"Depth: {current_state['depth']:.1f} ft, Formation: {current_state['formation']}")
        logger.info(Fore.YELLOW + f"Current ROP: {current_state['rop']:.1f} ft/hr" + Style.RESET_ALL)
        print(f"Current ROP: {current_state['rop']:.1f} ft/hr")

        logger.info(Fore.BLUE + "\nRunning parameter optimization..." + Style.RESET_ALL)
        print("\nRunning parameter optimization...")
        _, optimal_params = optimizer.optimize_parameters(current_state)

        logger.info(Fore.CYAN + "ML-Optimized Parameters:" + Style.RESET_ALL)
        print("ML-Optimized Parameters:")
        logger.info(Fore.CYAN + f"- WOB: {optimal_params['weight_on_bit']:.1f} kips" + Style.RESET_ALL)
        print(f"- WOB: {optimal_params['weight_on_bit']:.1f} kips")
        logger.info(Fore.CYAN + f"- RPM: {optimal_params['rotary_speed']:.1f} RPM" + Style.RESET_ALL)
        print(f"- RPM: {optimal_params['rotary_speed']:.1f} RPM")
        logger.info(Fore.CYAN + f"- Flow Rate: {optimal_params['flow_rate']:.1f} GPM" + Style.RESET_ALL)
        print(f"- Flow Rate: {optimal_params['flow_rate']:.1f} GPM")
        logger.info(Fore.CYAN + f"- Predicted ROP: {optimal_params['predicted_rop']:.1f} ft/hr" + Style.RESET_ALL)
        print(f"- Predicted ROP: {optimal_params['predicted_rop']:.1f} ft/hr")
        logger.info(Fore.CYAN + f"- Bit Failure Probability: {optimal_params['bit_failure_prob']:.2f}" + Style.RESET_ALL)
        print(f"- Bit Failure Probability: {optimal_params['bit_failure_prob']:.2f}")

        logger.info(Fore.BLUE + "\nGetting LLM-enhanced recommendation..." + Style.RESET_ALL)
        print("\nGetting LLM-enhanced recommendation...")
        recommendation = assistant.get_optimization_recommendation(optimal_params)
        logger.info(Fore.LIGHTWHITE_EX + "\nRecommendation:" + Style.RESET_ALL)
        print("\nRecommendation:")
        logger.info(Fore.LIGHTWHITE_EX + "-" * 40 + Style.RESET_ALL)
        print("-" * 40)
        logger.info(Fore.LIGHTGREEN_EX + recommendation + Style.RESET_ALL)
        print(recommendation)
        logger.info(Fore.LIGHTWHITE_EX + "-" * 40 + Style.RESET_ALL)
        print("-" * 40)

        if i < len(sim_indices) - 1:
            time.sleep(1)

    logger.info(Fore.GREEN + "\n5. Generating comprehensive drilling report with LLM..." + Style.RESET_ALL)
    print("\n5. Generating comprehensive drilling report with LLM...")
    report = assistant.generate_report()
    logger.info(Fore.LIGHTYELLOW_EX + "\nDRILLING REPORT:" + Style.RESET_ALL)
    print("\nDRILLING REPORT:")
    logger.info(Fore.LIGHTYELLOW_EX + "=" * 80 + Style.RESET_ALL)
    print("=" * 80)
    logger.info(Fore.LIGHTCYAN_EX + report + Style.RESET_ALL)
    print(report)
    logger.info(Fore.LIGHTYELLOW_EX + "=" * 80 + Style.RESET_ALL)
    print("=" * 80)

    logger.info(Fore.GREEN + "\nDemo completed successfully!" + Style.RESET_ALL)
    print("\nDemo completed successfully!")


def run_demo_and_save_to_word(filename="reports/drilling_demo_output.docx"):
    # Ensure the reports directory exists
    os.makedirs(os.path.dirname(filename), exist_ok=True)
    buffer = io.StringIO()
    sys_stdout = sys.stdout
    sys.stdout = buffer
    try:
        run_drilling_optimization_demo()
    finally:
        sys.stdout = sys_stdout
    output = buffer.getvalue()
    output = clean_text(output)
    doc = Document()
    # Insert image at the start
    image_path = os.path.join(os.path.dirname(__file__), '..', 'artifacts', '4.jpg')
    try:
        doc.add_picture(image_path, width=Inches(5))
    except Exception as e:
        print(f"Warning: Could not add image {image_path} to report: {e}")
    doc.add_heading('AI-Driven Drilling Optimization Demo Output', 0)
    format_markdown_to_word(doc, output)
    doc.save(filename)
    print(f"Demo output saved to {filename}")


if __name__ == "__main__":
    run_demo_and_save_to_word() 