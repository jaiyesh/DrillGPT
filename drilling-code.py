"""
Simplified AI-Driven Drilling Optimization with LLM Integration

This is a minimal implementation of an AI-driven drilling optimization system
that combines traditional machine learning with Large Language Models (LLMs).

Author: Jaiyesh Chahar
Date: May 2025
"""

import os
import json
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestRegressor, RandomForestClassifier
from sklearn.metrics import mean_squared_error, r2_score, classification_report
import datetime
import requests
import time
import io
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging
from colorama import Fore, Style, init as colorama_init
from pyfiglet import figlet_format
import re

# Set random seed for reproducibility
np.random.seed(42)

colorama_init(autoreset=True)

# Set up logger
logger = logging.getLogger("DrillGPT")
logger.setLevel(logging.INFO)
ch = logging.StreamHandler()
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
ch.setFormatter(formatter)
logger.handlers = []
logger.addHandler(ch)

def print_app_banner():
    banner = figlet_format("DrillGPT", font="slant")
    logger.info(Fore.CYAN + banner + Style.RESET_ALL)
    logger.info(Fore.YELLOW + "AI-Driven Drilling Optimization with LLM Integration" + Style.RESET_ALL)

#######################
# DATA GENERATION
#######################

def generate_synthetic_drilling_data(n_samples=500):
    """Generate synthetic drilling data with realistic relationships."""
    logger.info(Fore.LIGHTMAGENTA_EX + f"Generating {n_samples} synthetic drilling data points..." + Style.RESET_ALL)
    # Define formations and bit types
    formation_types = ['Sandstone', 'Shale', 'Limestone', 'Dolomite']
    bit_types = ['PDC', 'Tricone', 'Diamond']
    
    # Generate base parameters
    data = {
        'depth': np.cumsum(np.random.uniform(5, 30, n_samples)),
        'formation': np.random.choice(formation_types, n_samples),
        'bit_type': np.random.choice(bit_types, n_samples),
        'weight_on_bit': np.random.uniform(5, 30, n_samples),  # kips
        'rotary_speed': np.random.uniform(40, 200, n_samples),  # RPM
        'flow_rate': np.random.uniform(300, 1000, n_samples),  # GPM
        'bit_hours': np.zeros(n_samples),  # Hours on bit
    }
    
    df = pd.DataFrame(data)
    
    # Add formation hardness
    hardness_map = {'Sandstone': 5000, 'Shale': 3000, 'Limestone': 8000, 'Dolomite': 10000}
    df['formation_hardness'] = df['formation'].map(hardness_map)
    
    # Calculate bit hours
    for i in range(1, n_samples):
        if df.loc[i, 'bit_type'] != df.loc[i-1, 'bit_type']:
            df.loc[i, 'bit_hours'] = 0
        else:
            df.loc[i, 'bit_hours'] = df.loc[i-1, 'bit_hours'] + np.random.uniform(0.5, 2.0)
    
    # Calculate rate of penetration (ROP)
    # Base ROP with dependencies
    base_rop = 50 * np.random.normal(1, 0.2, n_samples)
    hardness_factor = 10000 / (df['formation_hardness'] + 1000)
    wob_factor = 0.8 + 0.4 * np.tanh((df['weight_on_bit'] - 15) / 5)
    rpm_factor = 0.6 + 0.8 * np.tanh((df['rotary_speed'] - 120) / 30)
    flow_factor = 0.7 + 0.6 * np.tanh((df['flow_rate'] - 650) / 150)
    bit_wear_factor = 1.2 - 0.5 * np.tanh((df['bit_hours'] - 50) / 20)
    
    df['rop'] = base_rop * hardness_factor * wob_factor * rpm_factor * flow_factor * bit_wear_factor
    df['rop'] = df['rop'].clip(5, 150)  # Reasonable bounds
    
    # Add vibration data
    df['vibration'] = 0.2 + 0.2 * np.random.normal(0, 1, n_samples) + 0.02 * df['weight_on_bit'] + 0.001 * df['rotary_speed']
    df['vibration'] = df['vibration'].clip(0, None)
    
    # Calculate bit failure probabilities
    failure_prob = 0.2 * np.tanh((df['bit_hours'] - 60) / 15) + 0.3 * np.tanh((df['vibration'] - 1.5) / 0.5)
    failure_prob = 0.2 + 0.8 * (failure_prob - failure_prob.min()) / (failure_prob.max() - failure_prob.min())
    df['bit_failure'] = (np.random.random(n_samples) < failure_prob).astype(int)
    
    return df

#######################
# ML MODELS
#######################

class DrillingModels:
    """Predictive models for drilling optimization."""
    
    def __init__(self):
        """Initialize models."""
        self.rop_model = None
        self.bit_failure_model = None
        self.feature_importances = None
        self.formation_types = ['Sandstone', 'Shale', 'Limestone', 'Dolomite']
        self.bit_types = ['PDC', 'Tricone', 'Diamond']
        self.train_columns = None  # Store training columns for consistent prediction
    
    def preprocess_data(self, df, target):
        """Preprocess drilling data for modeling."""
        # Create dummy variables for categorical features
        # We need to ensure all categories are present, even if not in this particular dataset
        
        # For formation
        formation_dummies = pd.get_dummies(df['formation'], prefix='formation')
        # Add missing formation columns if any
        for formation in self.formation_types:
            col_name = f'formation_{formation}'
            if col_name not in formation_dummies.columns:
                formation_dummies[col_name] = 0
                
        # For bit type
        bit_dummies = pd.get_dummies(df['bit_type'], prefix='bit')
        # Add missing bit columns if any
        for bit_type in self.bit_types:
            col_name = f'bit_{bit_type}'
            if col_name not in bit_dummies.columns:
                bit_dummies[col_name] = 0
        
        # Combine with original data
        data = pd.concat([df.drop(['formation', 'bit_type'], axis=1), formation_dummies, bit_dummies], axis=1)
        
        # Remove target from features if not needed
        if target != 'bit_failure' and 'bit_failure' in data.columns:
            data = data.drop('bit_failure', axis=1)
        
        # Split into X and y
        X = data.drop(target, axis=1) if target in data.columns else data
        y = data[target] if target in data.columns else None
        
        # If this is the first processing, store column order
        if self.train_columns is None and y is not None:  # Only store during training
            self.train_columns = X.columns.tolist()
        elif self.train_columns is not None:
            # Ensure columns match training data
            missing_cols = set(self.train_columns) - set(X.columns)
            for col in missing_cols:
                X[col] = 0
            # Ensure column order matches training data
            X = X[self.train_columns]
        
        return X, y
    
    def train_rop_model(self, df):
        """Train ROP prediction model."""
        logger.info(Fore.LIGHTBLUE_EX + "Training ROP prediction model..." + Style.RESET_ALL)
        X, y = self.preprocess_data(df, 'rop')
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
        logger.info(Fore.LIGHTBLUE_EX + f"Training samples: {len(X_train)}, Test samples: {len(X_test)}" + Style.RESET_ALL)
        self.rop_model = RandomForestRegressor(n_estimators=100, random_state=42)
        self.rop_model.fit(X_train, y_train)
        y_pred = self.rop_model.predict(X_test)
        mse = mean_squared_error(y_test, y_pred)
        r2 = r2_score(y_test, y_pred)
        logger.info(Fore.LIGHTGREEN_EX + f"ROP Model - MSE: {mse:.2f}, R²: {r2:.4f}" + Style.RESET_ALL)
        self.feature_importances = pd.Series(
            self.rop_model.feature_importances_, 
            index=X.columns
        ).sort_values(ascending=False)
        return {'mse': mse, 'r2': r2, 'X_test': X_test, 'y_test': y_test, 'y_pred': y_pred}
    
    def train_bit_failure_model(self, df):
        """Train bit failure prediction model."""
        logger.info(Fore.LIGHTBLUE_EX + "Training bit failure prediction model..." + Style.RESET_ALL)
        X, y = self.preprocess_data(df, 'bit_failure')
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42, stratify=y)
        logger.info(Fore.LIGHTBLUE_EX + f"Training samples: {len(X_train)}, Test samples: {len(X_test)}" + Style.RESET_ALL)
        self.bit_failure_model = RandomForestClassifier(n_estimators=100, random_state=42, class_weight='balanced')
        self.bit_failure_model.fit(X_train, y_train)
        y_pred = self.bit_failure_model.predict(X_test)
        logger.info(Fore.LIGHTGREEN_EX + "Bit Failure Model Results:" + Style.RESET_ALL)
        logger.info(Fore.LIGHTGREEN_EX + "\n" + classification_report(y_test, y_pred) + Style.RESET_ALL)
        return {'X_test': X_test, 'y_test': y_test, 'y_pred': y_pred}
    
    def predict_rop(self, new_data):
        """Predict ROP for new data."""
        if self.rop_model is None:
            raise ValueError("ROP model not trained")
        
        X, _ = self.preprocess_data(new_data, 'rop')
        return self.rop_model.predict(X)
    
    def predict_bit_failure_prob(self, new_data):
        """Predict bit failure probability for new data."""
        if self.bit_failure_model is None:
            raise ValueError("Bit failure model not trained")
        
        X, _ = self.preprocess_data(new_data, 'bit_failure')
        return self.bit_failure_model.predict_proba(X)[:, 1]

#######################
# PARAMETER OPTIMIZATION
#######################

class ParameterOptimizer:
    """Optimizer for drilling parameters."""
    
    def __init__(self, models):
        """Initialize the optimizer."""
        self.models = models
    
    def optimize_parameters(self, current_state):
        logger.info(Fore.LIGHTMAGENTA_EX + "Optimizing drilling parameters for current state..." + Style.RESET_ALL)
        logger.info(Fore.LIGHTMAGENTA_EX + f"Current WOB: {current_state.get('weight_on_bit', 'N/A')}, RPM: {current_state.get('rotary_speed', 'N/A')}, Flow Rate: {current_state.get('flow_rate', 'N/A')}" + Style.RESET_ALL)
        # Define parameter ranges to explore
        param_ranges = {
            'weight_on_bit': np.linspace(5, 30, 6),  # kips
            'rotary_speed': np.linspace(40, 200, 6),  # RPM
            'flow_rate': np.linspace(300, 1000, 6),  # GPM
        }
        
        # Test all parameter combinations
        results = []
        
        for wob in param_ranges['weight_on_bit']:
            for rpm in param_ranges['rotary_speed']:
                for flow in param_ranges['flow_rate']:
                    # Create test dataframe with current parameters
                    test_state = current_state.copy()
                    test_state['weight_on_bit'] = wob
                    test_state['rotary_speed'] = rpm
                    test_state['flow_rate'] = flow
                    
                    # Create a DataFrame for prediction
                    df_test = pd.DataFrame([test_state])
                    
                    # Use the models to predict ROP and bit failure probability
                    predicted_rop = self.models.predict_rop(df_test)[0]
                    bit_failure_prob = self.models.predict_bit_failure_prob(df_test)[0]
                    
                    # Store results
                    results.append({
                        'weight_on_bit': wob,
                        'rotary_speed': rpm,
                        'flow_rate': flow,
                        'predicted_rop': predicted_rop,
                        'bit_failure_prob': bit_failure_prob
                    })
        
        # Convert to DataFrame
        results_df = pd.DataFrame(results)
        
        # Find optimal parameters (max ROP with failure prob < 0.5)
        valid_results = results_df[results_df['bit_failure_prob'] < 0.5]
        if len(valid_results) > 0:
            optimal_row = valid_results.loc[valid_results['predicted_rop'].idxmax()]
        else:
            # If all combinations have high failure probability, just pick max ROP
            optimal_row = results_df.loc[results_df['predicted_rop'].idxmax()]
        
        return results_df, optimal_row


#######################
# LLM INTEGRATION
#######################

class LLMService:
    """Service for interacting with OpenAI's API."""
    
    @staticmethod
    def chat_completion(messages, model="gpt-4o", temperature=0.7, max_tokens=2000):
        """Make a chat completion request to OpenAI API."""
        api_key = os.environ.get("OPENAI_API_KEY", "your-api-key-here")
        
        # In a real application, add proper error handling and retries
        try:
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}"
                },
                json={
                    "model": model,
                    "messages": messages,
                    "temperature": temperature,
                    "max_tokens": max_tokens
                }
            )
            response.raise_for_status()
            return response.json()
        except:
            # For demo purposes, return a mock response if API call fails
            return LLMService._generate_mock_response(messages)
    
    @staticmethod
    def _generate_mock_response(messages):
        """Generate a mock response for demonstration purposes."""
        # Extract the last user message
        last_message = next((m for m in reversed(messages) if m["role"] == "user"), None)
        content = last_message.get("content", "") if last_message else ""
        
        # Simple rule-based mock responses
        if "report" in content.lower():
            mock_response = "Here's a drilling performance report based on the data:\n\n"
            mock_response += "- Average ROP: 45.3 ft/hr\n"
            mock_response += "- Maximum ROP achieved: 87.2 ft/hr\n"
            mock_response += "- Current bit wear: Moderate (43%)\n"
            mock_response += "- Risk assessment: Low to moderate risk of bit failure\n\n"
            mock_response += "Recommended actions:\n"
            mock_response += "1. Continue monitoring vibration levels\n"
            mock_response += "2. Consider increasing WOB by 2-3 kips\n"
            mock_response += "3. Maintain current RPM values"
        elif "optimize" in content.lower():
            mock_response = "Based on current drilling conditions, I recommend the following parameter adjustments:\n\n"
            mock_response += "- Increase WOB from 18.5 to 22.0 kips\n"
            mock_response += "- Decrease RPM from 150 to 130\n"
            mock_response += "- Maintain flow rate at 650 GPM\n\n"
            mock_response += "These adjustments should improve ROP by approximately 15-20% while keeping vibration levels within acceptable ranges."
        else:
            mock_response = "I'm your drilling assistant. I can help with generating reports, optimizing drilling parameters, explaining current conditions, or providing recommendations for improving performance."
        
        return {"choices": [{"message": {"content": mock_response}}]}

class DrillingAssistant:
    """LLM-enhanced drilling assistant."""
    
    def __init__(self, llm_service, drilling_models, current_state=None):
        """Initialize the drilling assistant."""
        self.llm_service = llm_service
        self.models = drilling_models
        self.current_state = current_state
        self.conversation_history = []
    
    def add_message(self, role, content):
        """Add a message to the conversation history."""
        self.conversation_history.append({"role": role, "content": content})
    
    def format_current_state(self):
        """Format the current drilling state as a text string."""
        if self.current_state is None:
            return "No current drilling state available."
        
        state = self.current_state
        
        result = "--- CURRENT DRILLING STATE ---\n"
        result += f"Depth: {state.get('depth', 'N/A'):.1f} ft\n"
        result += f"Formation: {state.get('formation', 'N/A')}\n"
        result += f"Bit Type: {state.get('bit_type', 'N/A')}\n"
        result += f"Bit Hours: {state.get('bit_hours', 'N/A'):.1f} hrs\n\n"
        
        result += "--- CURRENT PARAMETERS ---\n"
        result += f"Weight on Bit (WOB): {state.get('weight_on_bit', 'N/A'):.1f} kips\n"
        result += f"Rotary Speed (RPM): {state.get('rotary_speed', 'N/A'):.1f} RPM\n"
        result += f"Flow Rate: {state.get('flow_rate', 'N/A'):.1f} GPM\n\n"
        
        result += "--- KEY INDICATORS ---\n"
        result += f"ROP: {state.get('rop', 'N/A'):.1f} ft/hr\n"
        result += f"Vibration: {state.get('vibration', 'N/A'):.2f}\n"
        
        # Add ML predictions if available
        if self.models.rop_model is not None and self.models.bit_failure_model is not None:
            df = pd.DataFrame([state])
            predicted_rop = self.models.predict_rop(df)[0]
            bit_failure_prob = self.models.predict_bit_failure_prob(df)[0]
            
            result += f"Predicted ROP: {predicted_rop:.1f} ft/hr\n"
            result += f"Bit Failure Probability: {bit_failure_prob:.2f}\n"
        
        return result
    
    def query_llm(self, query):
        logger.info(Fore.LIGHTYELLOW_EX + f"Querying LLM: {query[:60]}..." + Style.RESET_ALL)
        # Prepare system prompt
        system_prompt = """
        You are DrillGPT, an AI drilling engineering assistant with expertise in drilling optimization, 
        well operations, and petroleum engineering. You help drilling engineers interpret data, optimize 
        parameters, troubleshoot issues, and make data-driven decisions. 
        
        You should always:
        1. Analyze the drilling data provided before responding
        2. Consider both data-driven insights and engineering principles
        3. Provide specific, actionable recommendations when possible
        4. Explain your reasoning clearly
        """
        
        # Prepare messages
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "system", "content": f"Current drilling context:\n\n{self.format_current_state()}"}
        ]
        
        # Add conversation history
        messages.extend(self.conversation_history)
        
        # Add the new query
        messages.append({"role": "user", "content": query})
        
        # Get response from LLM
        response = self.llm_service.chat_completion(messages)
        response_content = response["choices"][0]["message"]["content"]
        
        # Add the query and response to conversation history
        self.add_message("user", query)
        self.add_message("assistant", response_content)
        
        return response_content
    
    def generate_report(self):
        logger.info(Fore.LIGHTYELLOW_EX + "Generating drilling report using LLM..." + Style.RESET_ALL)
        return self.query_llm("Generate a comprehensive daily drilling report based on the current data.")
    
    def get_optimization_recommendation(self, optimal_params):
        logger.info(Fore.LIGHTYELLOW_EX + "Getting optimization recommendation from LLM..." + Style.RESET_ALL)
        current = self.current_state
        prompt = f"""
        The ML model recommends the following parameter changes:
        - WOB: {current.get('weight_on_bit', 'N/A'):.1f} → {optimal_params.get('weight_on_bit', 'N/A'):.1f} kips
        - RPM: {current.get('rotary_speed', 'N/A'):.1f} → {optimal_params.get('rotary_speed', 'N/A'):.1f} RPM
        - Flow Rate: {current.get('flow_rate', 'N/A'):.1f} → {optimal_params.get('flow_rate', 'N/A'):.1f} GPM
        
        Predicted ROP improvement: {optimal_params.get('predicted_rop', 0) - current.get('rop', 0):.1f} ft/hr
        
        Please explain why these changes would help and any precautions that should be taken.
        """
        return self.query_llm(prompt)

#######################
# WORKFLOW AND DEMO
#######################

def run_drilling_optimization_demo():
    """Run a demonstration of the LLM-enhanced drilling optimization system with colorful logging."""
    print_app_banner()
    logger.info(Fore.GREEN + "1. Generating synthetic drilling data..." + Style.RESET_ALL)
    print("\n1. Generating synthetic drilling data...")
    drilling_data = generate_synthetic_drilling_data(n_samples=500)
    logger.info(Fore.CYAN + f"   Generated {len(drilling_data)} data points" + Style.RESET_ALL)
    print(f"   Generated {len(drilling_data)} data points")

    logger.info(Fore.GREEN + "2. Training ML models..." + Style.RESET_ALL)
    print("\n2. Training ML models...")
    models = DrillingModels()
    models.train_rop_model(drilling_data)
    models.train_bit_failure_model(drilling_data)

    logger.info(Fore.GREEN + "3. Setting up parameter optimizer and drilling assistant..." + Style.RESET_ALL)
    print("\n3. Setting up parameter optimizer and drilling assistant...")
    optimizer = ParameterOptimizer(models)
    llm_service = LLMService()
    assistant = DrillingAssistant(llm_service, models)

    logger.info(Fore.GREEN + "4. Simulating real-time drilling optimization..." + Style.RESET_ALL)
    print("\n4. Simulating real-time drilling optimization...")
    sim_indices = [100, 200, 300]

    for i, idx in enumerate(sim_indices):
        logger.info(Fore.MAGENTA + f"\n--- Step {i+1}/{len(sim_indices)} ---" + Style.RESET_ALL)
        print(f"\n--- Step {i+1}/{len(sim_indices)} ---")
        current_state = drilling_data.iloc[idx].to_dict()
        assistant.current_state = current_state
        logger.info(Fore.YELLOW + f"Depth: {current_state['depth']:.1f} ft, Formation: {current_state['formation']}" + Style.RESET_ALL)
        print(f"Depth: {current_state['depth']:.1f} ft, Formation: {current_state['formation']}")
        logger.info(Fore.YELLOW + f"Current ROP: {current_state['rop']:.1f} ft/hr" + Style.RESET_ALL)
        print(f"Current ROP: {current_state['rop']:.1f} ft/hr")

        logger.info(Fore.BLUE + "\nRunning parameter optimization..." + Style.RESET_ALL)
        print("\nRunning parameter optimization...")
        _, optimal_params = optimizer.optimize_parameters(current_state)

        logger.info(Fore.CYAN + "ML-Optimized Parameters:" + Style.RESET_ALL)
        print("ML-Optimized Parameters:")
        logger.info(Fore.CYAN + f"- WOB: {optimal_params['weight_on_bit']:.1f} kips" + Style.RESET_ALL)
        print(f"- WOB: {optimal_params['weight_on_bit']:.1f} kips")
        logger.info(Fore.CYAN + f"- RPM: {optimal_params['rotary_speed']:.1f} RPM" + Style.RESET_ALL)
        print(f"- RPM: {optimal_params['rotary_speed']:.1f} RPM")
        logger.info(Fore.CYAN + f"- Flow Rate: {optimal_params['flow_rate']:.1f} GPM" + Style.RESET_ALL)
        print(f"- Flow Rate: {optimal_params['flow_rate']:.1f} GPM")
        logger.info(Fore.CYAN + f"- Predicted ROP: {optimal_params['predicted_rop']:.1f} ft/hr" + Style.RESET_ALL)
        print(f"- Predicted ROP: {optimal_params['predicted_rop']:.1f} ft/hr")
        logger.info(Fore.CYAN + f"- Bit Failure Probability: {optimal_params['bit_failure_prob']:.2f}" + Style.RESET_ALL)
        print(f"- Bit Failure Probability: {optimal_params['bit_failure_prob']:.2f}")

        logger.info(Fore.BLUE + "\nGetting LLM-enhanced recommendation..." + Style.RESET_ALL)
        print("\nGetting LLM-enhanced recommendation...")
        recommendation = assistant.get_optimization_recommendation(optimal_params)
        logger.info(Fore.LIGHTWHITE_EX + "\nRecommendation:" + Style.RESET_ALL)
        print("\nRecommendation:")
        logger.info(Fore.LIGHTWHITE_EX + "-" * 40 + Style.RESET_ALL)
        print("-" * 40)
        logger.info(Fore.LIGHTGREEN_EX + recommendation + Style.RESET_ALL)
        print(recommendation)
        logger.info(Fore.LIGHTWHITE_EX + "-" * 40 + Style.RESET_ALL)
        print("-" * 40)

        if i < len(sim_indices) - 1:
            time.sleep(1)

    logger.info(Fore.GREEN + "\n5. Generating comprehensive drilling report with LLM..." + Style.RESET_ALL)
    print("\n5. Generating comprehensive drilling report with LLM...")
    report = assistant.generate_report()
    logger.info(Fore.LIGHTYELLOW_EX + "\nDRILLING REPORT:" + Style.RESET_ALL)
    print("\nDRILLING REPORT:")
    logger.info(Fore.LIGHTYELLOW_EX + "=" * 80 + Style.RESET_ALL)
    print("=" * 80)
    logger.info(Fore.LIGHTCYAN_EX + report + Style.RESET_ALL)
    print(report)
    logger.info(Fore.LIGHTYELLOW_EX + "=" * 80 + Style.RESET_ALL)
    print("=" * 80)

    logger.info(Fore.GREEN + "\nDemo completed successfully!" + Style.RESET_ALL)
    print("\nDemo completed successfully!")

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    hr.append(bottom)
    p._p.get_or_add_pPr().append(hr)

def add_markdown_runs(para, text):
    """
    Add runs to a paragraph with bold and italic formatting for markdown.
    Supports **bold**, *italic*, and normal text.
    """
    import re
    # Pattern for **bold**, *italic*, and normal text
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    tokens = re.findall(pattern, text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = para.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = para.add_run(token[1:-1])
            run.italic = True
        else:
            para.add_run(token)

def format_markdown_to_word(doc, text):
    import re
    lines = text.split('\n')
    for line in lines:
        stripped = line.strip()
        # Headings
        if stripped.startswith('### '):
            para = doc.add_paragraph(stripped[4:], style='Heading 3')
        elif stripped.startswith('## '):
            para = doc.add_paragraph(stripped[3:], style='Heading 2')
        elif stripped.startswith('# '):
            para = doc.add_paragraph(stripped[2:], style='Heading 1')
        # Horizontal lines
        elif stripped.startswith('---') or stripped.startswith('==='):
            add_horizontal_line(doc)
        # Numbered list
        elif re.match(r'^\d+\.', stripped):
            para = doc.add_paragraph(style='List Number')
            add_markdown_runs(para, stripped)
        # Bullet list (handles both - and •)
        elif re.match(r'^[-*•]\s+', stripped):
            para = doc.add_paragraph(style='List Bullet')
            # Remove the bullet symbol and following space
            content = re.sub(r'^[-*•]\s+', '', stripped)
            add_markdown_runs(para, content)
        # Bold/italic inline in normal text
        elif '**' in stripped or '*' in stripped:
            para = doc.add_paragraph()
            add_markdown_runs(para, stripped)
        # Default
        else:
            doc.add_paragraph(line)

def clean_text(text):
    # Remove ANSI escape sequences (color codes)
    ansi_escape = re.compile(r'\x1B[@-_][0-?]*[ -/]*[@-~]')
    text = ansi_escape.sub('', text)
    # Remove NULL bytes and other non-printable/control characters except newlines and tabs
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    return text

def run_demo_and_save_to_word(filename="drilling_demo_output.docx"):
    logger.info(Fore.LIGHTMAGENTA_EX + f"Running demo and saving output to Word file: {filename}" + Style.RESET_ALL)
    buffer = io.StringIO()
    sys_stdout = sys.stdout
    sys.stdout = buffer
    try:
        run_drilling_optimization_demo()
    finally:
        sys.stdout = sys_stdout
    output = buffer.getvalue()
    output = clean_text(output)  # Clean the output before saving to Word
    doc = Document()
    doc.add_heading('AI-Driven Drilling Optimization Demo Output', 0)
    format_markdown_to_word(doc, output)
    doc.save(filename)
    logger.info(Fore.LIGHTGREEN_EX + f"Demo output saved to {filename}" + Style.RESET_ALL)

# Run the demo if executing this script directly
if __name__ == "__main__":
    # run_drilling_optimization_demo()  # Comment this out to avoid double printing
    run_demo_and_save_to_word()
