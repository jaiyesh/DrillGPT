"""Utilities for formatting and saving reports (Word, Markdown, etc.)."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

# --- Formatting and Report Utilities ---

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    hr.append(bottom)
    p._p.get_or_add_pPr().append(hr)

def add_markdown_runs(para, text):
    """
    Add runs to a paragraph with bold and italic formatting for markdown.
    Supports **bold**, *italic*, and normal text.
    """
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    tokens = re.findall(pattern, text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = para.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = para.add_run(token[1:-1])
            run.italic = True
        else:
            para.add_run(token)

def format_markdown_to_word(doc, text):
    lines = text.split('\n')
    for line in lines:
        stripped = line.strip()
        # Headings
        if stripped.startswith('### '):
            para = doc.add_paragraph(stripped[4:], style='Heading 3')
        elif stripped.startswith('## '):
            para = doc.add_paragraph(stripped[3:], style='Heading 2')
        elif stripped.startswith('# '):
            para = doc.add_paragraph(stripped[2:], style='Heading 1')
        # Horizontal lines
        elif stripped.startswith('---') or stripped.startswith('==='):
            add_horizontal_line(doc)
        # Numbered list
        elif re.match(r'^\d+\.', stripped):
            para = doc.add_paragraph(style='List Number')
            add_markdown_runs(para, stripped)
        # Bullet list (handles both - and •)
        elif re.match(r'^[-*•]\s+', stripped):
            para = doc.add_paragraph(style='List Bullet')
            # Remove the bullet symbol and following space
            content = re.sub(r'^[-*•]\s+', '', stripped)
            add_markdown_runs(para, content)
        # Bold/italic inline in normal text
        elif '**' in stripped or '*' in stripped:
            para = doc.add_paragraph()
            add_markdown_runs(para, stripped)
        # Default
        else:
            doc.add_paragraph(line)

def clean_text(text):
    # Remove ANSI escape sequences (color codes)
    ansi_escape = re.compile(r'\x1B[@-_][0-?]*[ -/]*[@-~]')
    text = ansi_escape.sub('', text)
    # Remove NULL bytes and other non-printable/control characters except newlines and tabs
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    return text 